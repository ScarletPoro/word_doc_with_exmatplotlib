from matplotlib import pyplot
from docx import Document
from docx.shared import Inches
import os

# Funzione per sostituire un placeholder con un'immagine
def replace_placeholder_with_image(doc, placeholder, image_path):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Sostituisci il placeholder con l'immagine
            paragraph.clear()  # Cancella il testo del paragrafo
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(5))  # Aggiunge il grafico come immagine

# Funzione per creare e salvare un grafico a linee
def crea_grafico_linee(x, y1, y2, titolo, xlabel, ylabel, legend1, legend2, output_path):
    # Imposta lo stile del grafico
    pyplot.style.use('dark_background')

    # Creazione del grafico a linee
    pyplot.plot(x, y1, label=legend1, color="red", marker='o', linewidth=2)
    pyplot.plot(x, y2, label=legend2, color="blue", linestyle='dotted', marker='.')

    # Aggiunta di titoli e legenda
    pyplot.title(titolo)
    pyplot.xlabel(xlabel)
    pyplot.ylabel(ylabel)
    pyplot.legend()
    pyplot.grid()

    # Salva il grafico come immagine PNG
    pyplot.savefig(output_path)
    pyplot.clf()  # Pulisci il grafico per evitare sovrapposizioni future

# Funzione per creare e salvare un grafico a barre orizzontali
def crea_grafico_barre(x, y, titolo, xlabel, ylabel, output_path):
    pyplot.style.use('ggplot')

    # Creazione del grafico a barre orizzontali
    pyplot.barh(x, y, color='green')

    # Aggiunta di titoli e legenda
    pyplot.title(titolo)
    pyplot.xlabel(xlabel)
    pyplot.ylabel(ylabel)

    # Salva il grafico come immagine PNG
    pyplot.savefig(output_path)
    pyplot.clf()  # Pulisci il grafico

# Percorso del file Word esistente
word_file_path = "documento_esistente.docx"

# Apri il documento esistente
doc = Document(word_file_path)

# Crea il primo grafico a linee e ottieni il percorso dell'immagine
image_path_1 = "grafico1.png"
x = [25, 30, 32, 50, 75]
y1 = [100, 82, 50, 33, 92]
y2 = [90, 70, 150, 70, 22]
crea_grafico_linee(x, y1, y2, 'Grafico 1 - Andamento pesi', 'x - età', 'y - peso', 'italiani', 'francesi', image_path_1)

# Sostituisci il placeholder <<DIAGRAMMA1>> con il primo grafico a linee
replace_placeholder_with_image(doc, '<<DIAGRAMMA1>>', image_path_1)

# Crea il secondo grafico a barre orizzontali e ottieni il percorso dell'immagine
image_path_2 = "grafico2.png"
y3 = [55, 65, 85, 95, 105]
x_labels = ['Gruppo A', 'Gruppo B', 'Gruppo C', 'Gruppo D', 'Gruppo E']
crea_grafico_barre(x_labels, y3, 'Grafico 2 - Confronto altezze', 'Altezza', 'Gruppi', image_path_2)

# Sostituisci il placeholder <<DIAGRAMMA2>> con il secondo grafico a barre
replace_placeholder_with_image(doc, '<<DIAGRAMMA2>>', image_path_2)

# Salva il documento aggiornato
doc.save("documento_aggiornato.docx")

# Pulisci le immagini temporanee
os.remove(image_path_1)
os.remove(image_path_2)

print("Placeholder sostituiti e documento salvato come 'documento_aggiornato.docx'")
